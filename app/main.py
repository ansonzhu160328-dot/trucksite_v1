from fastapi import FastAPI, Request, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from docx import Document
import tempfile
import os
from pathlib import Path

import base64
import shutil
import subprocess
from datetime import datetime
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


from app.schemas import CalcRequest
from app.calc import calc_plan

app = FastAPI(title="Truck Charging Site V1", version="0.3.0")

app.mount("/static", StaticFiles(directory="static"), name="static")

BASE_DIR = Path(__file__).resolve().parent.parent
PRODUCT_ASSETS_DIR = BASE_DIR / "assets" / "product"
ALLOWED_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp"}
FINANCE_TEXT = """
一、融资方案概述
本方案可针对场站建设及设备采购提供整站融资支持，参考年化成本约为6厘。融资范围可覆盖充电站项目整体投入，适用于具备一定经营基础、信用记录良好的企业客户。

二、准入条件
1.物流公司实际经营满三年，或老公司实际经营满三年；如项目公司设立时间较短，可由老公司提供担保。对于运营商客户，可适当放宽年限要求。
2.业主企业及实际控制人当前征信无逾期，不存在影响经营的重大诉讼、失信被执行或限制高消费情况，且当前无与金融机构借款纠纷。
3.可接受有限公司、股份公司、个体工商户、合伙企业、国有参股企业准入，国有控股企业暂不介入。
4.近六个月月末结余较为平稳，可根据前半年每月结余情况核定融资额度，接受未开票的私户经营收入作为辅助参考。

三、目标客户
1.已具备一定经验的充电场站运营商；
2.物流公司或重卡车队客户；
3.新进入充电行业，但此前已有其他稳定经营业务的公司。

四、还款方式及担保措施
1.还款方式：按月等额还款，融资期限一般为2至5年。
2.担保措施：原则上需两人或以上提供担保，担保人合计持股比例建议超过70%。

五、方案特点
1.不限设备类型：新购充电桩设备、换电站设备、高压设备、电池及储能设备等均可办理分期融资。
2.可提前起租：可在设备发货前支付30%-50%款项，设备进场并由我司拍照确认后支付余款。
3.额度上限较高：融资额度一般为30万元至2000万元。
4.不上征信，不影响企业后续银行贷款及授信。
5.不限区域，可面向全国开展业务。
6.可支持全额融资，最高可覆盖合同金额的100%。
7.0手续费、0保证金。

六、操作流程
1.锁定意向：设备厂商与客户确认分期采购意向。
2.资料初审：客户提供流水、征信、财务报表等基础资料。
3.项目尽调及审批：开展现场尽调，补充所需资料，并提交风控审批。
4.签约放款：授信审批通过后完成签约，设备供应商向我司开票（若客户需发票，则按约定处理）；设备发货前我司支付部分设备款，待我司资产部门验收后支付余款。

七、资料清单
1.基本资料：营业执照、公司章程、征信报告。
2.财务资料：内部财务报表。
3.资产资料：场地租赁合同、设备采购合同、下游合同（如有）、银行流水。
4.担保人资料：身份证、房产证、银行流水、征信报告。

八、项目案例（以100万元、5年期为例）
1.承租人：某有限公司
2.租赁类型：直租
3.租赁物：充换电站系统（整个项目工程）
4.融资金额：100万元
5.融资比例：100%
6.租赁期限：5年
7.年化利率：5.31%（不含税）
8.每期租金：19174.04元
9.本息总额：130万元
10.还款方式：等额本息，按月还款
11.所有权安排：租赁期满后以1元形式转让设备所有权

九、特别说明
以上信息仅供参考，不构成任何承诺，具体融资方案及合同条款以最终签署文件为准。
""".strip()


@app.get("/")
def home():
    return FileResponse("static/index.html")


@app.post("/api/calculate")
def calculate(req: CalcRequest):
    data = req.model_dump()
    print("DEBUG /api/calculate keys:", sorted(list(data.keys())))
    result = calc_plan(data)
    return result


def build_report_doc(raw_data: dict) -> Document:
    req = CalcRequest.model_validate(raw_data)
    data = req.model_dump()
    result = calc_plan(data)

    # ===== 生成 Word =====
    doc = Document()

    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION_START
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # ===== 封皮标题 =====
    site_location = data.get("site_location", "").strip()

    if not site_location:
        site_location = "未填写场站位置"

    title_text = f"{site_location}重卡充电站初步设计方案"

    # =========================
    # 通用：字体 + 段落格式
    # =========================
    INDENT_2CH = Pt(28)  # 首行缩进 2 字符（宋体14号下约等于 28pt）
    LINE_SPACING = 1.5   # 1.5倍行距
    BULLET = "■"         # 统一条目标识符号

    def set_cn_font(run, size_pt=14, bold=False, font_name="宋体"):
        """
        中文：宋体
        英文/数字：Times New Roman
        """
        run.font.size = Pt(size_pt)
        run.bold = bold

        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()

        # 中文字体
        rFonts.set(qn('w:eastAsia'), font_name)

        # 英文 & 数字字体
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')


    def format_para(p, first_line_indent=False):
        p.paragraph_format.line_spacing = LINE_SPACING
        if first_line_indent:
            p.paragraph_format.first_line_indent = INDENT_2CH

    def add_cover_line(text, size_pt=14, bold=False, align_center=True):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_cn_font(run, size_pt=size_pt, bold=bold, font_name="宋体")
        format_para(p, first_line_indent=False)

    def add_title(text):
        # 一级标题：宋体14pt加粗，不缩进，1.5倍行距
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_cn_font(run, size_pt=14, bold=True, font_name="宋体")
        format_para(p, first_line_indent=False)

    def add_body(text):
        # 正文：宋体14pt不加粗，首行缩进2字符，1.5倍行距
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_cn_font(run, size_pt=14, bold=False, font_name="宋体")
        format_para(p, first_line_indent=True)

    def add_finance_body(text):
        # 金融附件正文：宋体14pt，单倍行距，段前段后0磅，首行缩进2字符
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_cn_font(run, size_pt=14, bold=False, font_name="宋体")
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = INDENT_2CH

    def add_item(text):
        # 条目：用 ■ 符号；不做首行缩进（避免符号被挤歪），1.5倍行距
        p = doc.add_paragraph()
        run = p.add_run(f"{BULLET} {text}")
        set_cn_font(run, size_pt=14, bold=False, font_name="宋体")
        format_para(p, first_line_indent=False)

    def add_blank_line():
        # 章节结束空一行
        p = doc.add_paragraph("")
        p.paragraph_format.line_spacing = LINE_SPACING

    def format_report_table(table):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        tbl = table._tbl
        tbl_pr = tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            elem = OxmlElement(f'w:{edge}')
            elem.set(qn('w:val'), 'single')
            elem.set(qn('w:sz'), '8')
            elem.set(qn('w:space'), '0')
            elem.set(qn('w:color'), '000000')
            borders.append(elem)
        tbl_pr.append(borders)

        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1

    def add_simple_table(headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        for i, text in enumerate(headers):
            cell_p = table.rows[0].cells[i].paragraphs[0]
            cell_p.paragraph_format.line_spacing = LINE_SPACING
            run = cell_p.add_run(str(text))
            set_cn_font(run, size_pt=14, bold=True, font_name="宋体")
        for row in rows:
            cells = table.add_row().cells
            for i, text in enumerate(row):
                cell_p = cells[i].paragraphs[0]
                run = cell_p.add_run(str(text))
                set_cn_font(run, size_pt=14, bold=False, font_name="宋体")

        format_report_table(table)

    def add_body_bold(text, first_line_indent=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_cn_font(run, size_pt=14, bold=True, font_name="宋体")
        format_para(p, first_line_indent=first_line_indent)

    def add_numbered(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_cn_font(run, size_pt=14, bold=False, font_name="宋体")
        format_para(p, first_line_indent=False)

    def _hide_table_borders(table):
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            elem = OxmlElement(f'w:{edge}')
            elem.set(qn('w:val'), 'nil')
            borders.append(elem)
        tbl_pr.append(borders)

    def add_cover_page():
        section1 = doc.sections[0]
        usable_height = section1.page_height - section1.top_margin - section1.bottom_margin

        table = doc.add_table(rows=3, cols=1)
        _hide_table_borders(table)

        row_heights = [int(usable_height * 0.3), int(usable_height * 0.4), int(usable_height * 0.3)]
        for idx, row in enumerate(table.rows):
            row.height = row_heights[idx]
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        # 中间：标题（水平+垂直居中）
        mid_cell = table.cell(1, 0)
        mid_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        mid_para = mid_cell.paragraphs[0]
        mid_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = mid_para.add_run(title_text)
        set_cn_font(run, size_pt=22, bold=True, font_name="宋体")
        format_para(mid_para, first_line_indent=False)

        # 底部：编制单位/日期（左对齐+底部对齐）
        bottom_cell = table.cell(2, 0)
        bottom_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

        p1 = bottom_cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run("编制单位：广东盈通智联数字技术有限公司")
        set_cn_font(r1, size_pt=14, bold=False, font_name="宋体")
        format_para(p1, first_line_indent=False)

        p2 = bottom_cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}")
        set_cn_font(r2, size_pt=14, bold=False, font_name="宋体")
        format_para(p2, first_line_indent=False)

    def _set_section_page_start(section, start_num=1):
        sect_pr = section._sectPr
        for child in list(sect_pr):
            if child.tag == qn('w:pgNumType'):
                sect_pr.remove(child)
        pg_num_type = OxmlElement('w:pgNumType')
        pg_num_type.set(qn('w:start'), str(start_num))
        sect_pr.append(pg_num_type)


    def _configure_body_header(section):
        section.header.is_linked_to_previous = False
        section.different_first_page_header_footer = False

        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.clear()
        run = p.add_run("广东盈通智联数字技术有限公司")
        set_cn_font(run, size_pt=14, bold=False, font_name="宋体")
        format_para(p, first_line_indent=False)

    def _add_footer_page_field(section):
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')

        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '

        fld_separate = OxmlElement('w:fldChar')
        fld_separate.set(qn('w:fldCharType'), 'separate')

        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_separate)
        run._r.append(fld_end)
    
    def normalize_attachments_selected(value):
        if isinstance(value, list):
            items = value
        elif isinstance(value, str):
            items = [value]
        else:
            items = []

        normalized = []
        for item in items:
            if not isinstance(item, str):
                continue
            v = item.strip().lower()
            if v in {"layout", "product", "finance"} and v not in normalized:
                normalized.append(v)
        return normalized

    def add_attach_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_cn_font(run, size_pt=14, bold=True, font_name="宋体")
        format_para(p, first_line_indent=False)

    def add_attach_hint(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_cn_font(run, size_pt=14, bold=False, font_name="宋体")
        format_para(p, first_line_indent=False)

    def parse_layout_png_data_url(layout_png_data_url: str):
        if not layout_png_data_url:
            return None
        b64 = layout_png_data_url
        if "base64," in layout_png_data_url:
            b64 = layout_png_data_url.split("base64,", 1)[1]
        try:
            return base64.b64decode(b64)
        except Exception as e:
            print("WARN layout_png_data_url decode failed:", e)
            return None

    def append_layout_attachment(data_dict: dict, attach_title: str):
        add_attach_title(attach_title)

        layout_png_data_url = (data_dict.get("layout_png_data_url") or "").strip()
        img_bytes = parse_layout_png_data_url(layout_png_data_url)
        if not img_bytes:
            add_attach_hint("（未获取到布局图）")
            return

        tmp_png = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        try:
            tmp_png.write(img_bytes)
            tmp_png.close()
            doc.add_picture(tmp_png.name, width=Cm(15))
        except Exception as e:
            print("WARN append layout image failed:", e)
            add_attach_hint("（未获取到布局图）")
        finally:
            try:
                os.remove(tmp_png.name)
            except Exception:
                pass

    def append_product_attachment(attach_title: str):
        add_attach_title(attach_title)

        if not PRODUCT_ASSETS_DIR.exists() or not PRODUCT_ASSETS_DIR.is_dir():
            add_attach_hint("（未配置产品图片）")
            return

        image_files = [
            p for p in PRODUCT_ASSETS_DIR.glob("*")
            if p.is_file() and p.suffix.lower() in ALLOWED_IMAGE_EXTS
        ]
        image_files.sort(key=lambda x: x.name)

        if not image_files:
            add_attach_hint("（未配置产品图片）")
            return

        inserted = False
        for image_path in image_files:
            try:
                if image_path.suffix.lower() == ".webp":
                    print("INFO skip unsupported webp for python-docx:", str(image_path))
                    continue
                doc.add_picture(str(image_path), width=Cm(15))
                doc.add_paragraph("")
                inserted = True
            except Exception as e:
                print("WARN append product image failed:", str(image_path), e)

        if not inserted:
            add_attach_hint("（未配置产品图片）")

    def append_finance_attachment(attach_title: str):
        add_attach_title(attach_title)

        text = (FINANCE_TEXT or "").strip()
        if not text:
            add_attach_hint("（未配置金融方案文本）")
            return

        for line in [x.strip() for x in text.splitlines() if x.strip()]:
            add_finance_body(line)


    # =========================
    # 封皮页（第1页）+ 分节（正文从第2页开始）
    # =========================
    add_cover_page()

    section2 = doc.add_section(WD_SECTION_START.NEW_PAGE)

    section2.footer.is_linked_to_previous = False
    section2.header.is_linked_to_previous = False
    _configure_body_header(section2)
    _set_section_page_start(section2, start_num=1)
    _add_footer_page_field(section2)


    # =========================
    # 一、项目基本情况
    # =========================
    add_title("一、项目基本情况")
    add_body("本项目拟建设重卡充电站1座，场站基本情况如下：")

    area = float(data.get('site_length_m', 0) or 0) * float(data.get('site_width_m', 0) or 0)

    add_simple_table(
        ["项目", "参数"],
        [
            ["场站位置", data.get('site_location', '')],
            ["场地长度", f"{data.get('site_length_m', 0)} m"],
            ["场地宽度", f"{data.get('site_width_m', 0)} m"],
            ["场地面积", f"{round(area, 2)} ㎡"],
            ["场地租金", f"{data.get('rent_yuan_per_sqm_month', 0)} 元/月"],
        ],
    )

    add_body("场地面积按场地长度与宽度计算得到，为后续设备布置及投资测算的重要依据。")
    add_blank_line()

    # =========================
    # 二、场站建设初步方案
    # =========================
    add_title("二、场站建设初步方案")
    add_body("根据场地条件、重卡充电需求以及设备功率配置，初步建议建设方案如下：")

    device_count = int(result.get('n_recommend', 0) or 0)
    gun_count = device_count * 2

    add_simple_table(
        ["项目", "参数"],
        [
            ["充电设备", "400kW一体机"],
            ["设备数量", f"{device_count}台"],
            ["充电枪数量", f"{gun_count}把"],
            ["配套变压器容量", f"{int(result.get('power_capacity_kva', 0) or 0)}kVA"],
            ["配套充电车位", f"{int(result.get('stalls', 0) or 0)}个"],
        ],
    )

    add_body("同时预留车辆通行及设备维护空间，以保证场站运行效率及安全性。")
    add_body("场站布局示意图详见附件1。")
    add_blank_line()

    # =========================
    # 三、项目投资估算
    # =========================
    add_title("三、项目投资估算")

    total_invest = round(float(result.get('invest_total_yuan', 0) or 0) / 10000, 2)
    power_invest = round(float(result.get('invest_power_yuan', 0) or 0) / 10000, 2)
    civil_invest = round(float(result.get('invest_civil_yuan', 0) or 0) / 10000, 2)
    equipment_invest = round(float(result.get('invest_pile_yuan', 0) or 0) / 10000, 2)

    add_body(f"根据当前建设方案，对场站建设投资进行初步测算，预计总投资约 {total_invest}万元，投资构成如下：")

    add_simple_table(
        ["投资项目", "金额"],
        [
            ["电力增容", f"{power_invest}万元"],
            ["场地土建", f"{civil_invest}万元"],
            ["充电设备", f"{equipment_invest}万元"],
            ["合计投资", f"{total_invest}万元"],
        ],
    )

    add_body("以上投资为初步估算，具体金额需根据实际工程实施情况进行调整。")
    add_blank_line()

    # =========================
    # 四、运营收益测算
    # =========================
    add_title("四、运营收益测算")
    add_body("在常规运营条件下，对场站经营收益进行初步测算：")

    annual_revenue = round(float(result.get('revenue_year_yuan', 0) or 0) / 10000, 2)
    annual_rent = round(float(result.get('rent_year_yuan', 0) or 0) / 10000, 2)
    labor_cost = round(float(result.get('labor_year_yuan', 0) or 0) / 10000, 2)
    net_cashflow = round(float(result.get('revenue_net_year_yuan', 0) or 0) / 10000, 2)
    payback = result.get('payback_net_years', None)
    payback_text = f"{round(float(payback), 2)}" if payback is not None else "N/A"

    add_simple_table(
        ["指标", "数值"],
        [
            ["年充电量", f"{int(result.get('energy_year_kwh', 0) or 0)}kWh"],
            ["服务费", f"{round(float(data.get('service_fee_yuan_per_kwh', 0) or 0), 2)}元/kWh"],
            ["年收入", f"{annual_revenue}万元"],
            ["年租金", f"{annual_rent}万元"],
            ["人工费用", f"{labor_cost}万元"],
            ["净年现金流", f"{net_cashflow}万元"],
            ["投资回收期", f"{payback_text}年"],
        ],
    )

    add_body_bold("测算假设条件", first_line_indent=False)
    add_body("单枪充电量：1000度/枪/天")
    add_body("充电服务费：0.3元/度")
    add_body("运营天数：330天/年")
    add_body("上述条件为典型运营场景下的参考参数，实际运营情况可能根据区域市场及车流情况有所变化。")
    add_blank_line()

    # =========================
    # 五、敏感性分析
    # =========================
    add_title("五、敏感性分析")
    add_body("为了评估关键参数变化对项目收益的影响，对充电量、服务费等因素进行组合分析（共27组）。")
    add_body("主要结论如下：")

    add_simple_table(
        ["情况", "年净收入", "投资回收期"],
        [
            ["最优情况", "85.54万元", "1.27年"],
            ["常规情况", "59.4万元", "1.83年"],
            ["最差情况", "28.51万元", "3.81年"],
        ],
    )

    add_body("分析结果表明，充电量及服务费为影响项目收益的主要因素，场站选址及客户资源对项目运营具有重要影响。")
    add_body("详细分析结果见附件相关数据表。")
    add_blank_line()

    # =========================
    # 六、结论与建议
    # =========================
    add_title("六、结论与建议")

    add_numbered(
        f"1、在常规运营条件下，本项目预计总投资约 {total_invest}万元，年净现金流约 {net_cashflow}万元，静态投资回收期约 {payback_text}年，项目整体投资收益较好。"
    )
    add_numbered("2、在最不利情况下，项目回收期约 3.81年，仍处于可接受范围。")
    add_numbered("3、建议优先选择物流车队密集区域建设，以保障充电利用率，提高项目运营收益。")

    add_blank_line()

    # ===== 文末附件：按前端选择动态插入（编号连续重排） =====
    attachments_selected = normalize_attachments_selected(raw_data.get("attachments_selected", []))

    attachment_defs = [
        ("layout", "场站布局示意图"),
        ("product", "产品及典型案例"),
        ("finance", "金融合作方案"),
    ]
    selected_attachments = [
        (kind, title) for kind, title in attachment_defs
        if kind in attachments_selected
    ]

    for idx, (kind, title) in enumerate(selected_attachments, start=1):
        if idx == 1 or kind in {"layout", "product"}:
            doc.add_page_break()
        attach_title = f"附件{idx}：{title}"
        if kind == "layout":
            append_layout_attachment(raw_data, attach_title)
        elif kind == "product":
            append_product_attachment(attach_title)
        elif kind == "finance":
            append_finance_attachment(attach_title)

    return doc


@app.post("/api/report_word")
async def report_word(request: Request):
    raw_data = await request.json()
    if not isinstance(raw_data, dict):
        raw_data = {}

    print("DEBUG /api/report_word keys:", sorted(list(raw_data.keys())))
    doc = build_report_doc(raw_data)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()

    return FileResponse(
        tmp.name,
        filename="trucksite_preliminary_design.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.post("/api/report_pdf")
async def report_pdf(req: CalcRequest, request: Request):
    data = req.model_dump()

    raw_data = await request.json()
    if not isinstance(raw_data, dict):
        raw_data = {}

    merged_data = dict(data)
    merged_data.update(raw_data)

    print("DEBUG /api/report_pdf keys:", sorted(list(merged_data.keys())))
    doc = build_report_doc(merged_data)

    # 依赖 LibreOffice（soffice）进行 headless 转换：
    # Ubuntu 安装：
    #   sudo apt update
    #   sudo apt install -y libreoffice
    #   soffice --version
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        docx_path = tmpdir_path / "report.docx"
        pdf_path = tmpdir_path / "report.pdf"
        doc.save(str(docx_path))

        cmd = [
            "soffice", "--headless", "--nologo", "--nofirststartwizard",
            "--convert-to", "pdf", "--outdir", str(tmpdir_path), str(docx_path)
        ]

        try:
            proc = subprocess.run(cmd, capture_output=True, text=True, check=False)
        except FileNotFoundError:
            raise HTTPException(status_code=500, detail="LibreOffice/soffice 未安装")

        if proc.returncode != 0:
            err = ((proc.stderr or proc.stdout or "soffice convert failed")[:1000]).strip()
            raise HTTPException(status_code=500, detail=f"PDF转换失败: {err}")

        if not pdf_path.exists() or pdf_path.stat().st_size <= 0:
            raise HTTPException(status_code=500, detail="PDF转换失败: 输出文件不存在或为空")

        # 方案2：先在临时目录转换，再拷贝到 delete=False 临时文件返回
        # 这样 TemporaryDirectory 可安全清理，同时 FileResponse 仍有稳定文件可读
        out_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        out_pdf.close()
        shutil.copyfile(str(pdf_path), out_pdf.name)

    return FileResponse(
        out_pdf.name,
        filename="trucksite_preliminary_design.pdf",
        media_type="application/pdf",
    )
